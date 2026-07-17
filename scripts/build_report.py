"""Build a review DOCX from the maintained course-design report source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parent.parent
SOURCE_MD = ROOT / "report-src" / "电子技术课程设计报告.md"
OUTPUT_DOCX = ROOT / "output" / "电子技术课程设计报告.docx"

FONT_BODY = "宋体"
FONT_HEI = "黑体"
FONT_LATIN = "Times New Roman"
FONT_MATH = "Cambria Math"
COLOR_BLACK = RGBColor(0, 0, 0)

BODY_SIZE = 12
TABLE_SIZE = 10.5
WORD_MATH_XSL = Path("/Applications/Microsoft Word.app/Contents/Resources/mathml2omml.xsl")
MATHML_NS = "http://www.w3.org/1998/Math/MathML"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_MATH_XSLT = None


def _mi(text, *, normal=False):
    variant = ' mathvariant="normal"' if normal else ""
    return f"<mi{variant}>{text}</mi>"


def _mn(text):
    return f"<mn>{text}</mn>"


def _mo(text):
    return f"<mo>{text}</mo>"


def _mtext(text):
    return f"<mtext>{text}</mtext>"


def _mrow(*parts):
    return f"<mrow>{''.join(parts)}</mrow>"


def _msub(base, subscript):
    return f"<msub>{base}{subscript}</msub>"


def _mfrac(numerator, denominator):
    return f"<mfrac>{numerator}{denominator}</mfrac>"


def _mover(expression):
    return f'<mover accent="true">{expression}<mo>¯</mo></mover>'


def _math(*parts):
    return f'<math xmlns="{MATHML_NS}">{_mrow(*parts)}</math>'


def _variable(name, subscript=None, *, subscript_normal=True):
    base = _mi(name)
    if subscript is None:
        return base
    sub = _mi(subscript, normal=subscript_normal) if not subscript.isdigit() else _mn(subscript)
    return _msub(base, sub)


def _signal(name):
    return _mi(name, normal=True)


def _unit(name):
    return _mtext(f" {name}")


_EM_SPACE = _mtext(" ")
MATHML_EXPRESSIONS = {
    r"t_H": _math(_variable("t", "H")),
    r"t_L": _math(_variable("t", "L")),
    r"f": _math(_variable("f")),
    r"R": _math(_variable("R")),
    r"\tau": _math(_variable("τ")),
    r"R_A": _math(_variable("R", "A")),
    r"R_B": _math(_variable("R", "B")),
    r"R_P": _math(_variable("R", "P")),
    r"C": _math(_variable("C")),
    r"R_B=4.7\ \mathrm{k\Omega}+R_P": _math(
        _variable("R", "B"), _mo("="), _mn("4.7"), _unit("kΩ"), _mo("+"), _variable("R", "P")
    ),
    r"f=1000\ \mathrm{Hz}": _math(_variable("f"), _mo("="), _mn("1000"), _unit("Hz")),
    r"R_A=4.7\ \mathrm{k\Omega}": _math(
        _variable("R", "A"), _mo("="), _mn("4.7"), _unit("kΩ")
    ),
    r"C=0.01\ \mathrm{\mu F}": _math(_variable("C"), _mo("="), _mn("0.01"), _unit("μF")),
    r"S_i=\overline{Y_i}": _math(
        _variable("S", "i", subscript_normal=False),
        _mo("="),
        _mover(_variable("Y", "i", subscript_normal=False)),
    ),
    r"S_0": _math(_variable("S", "0")),
    r"S_1": _math(_variable("S", "1")),
    r"S_2": _math(_variable("S", "2")),
    r"S_3": _math(_variable("S", "3")),
    r"t_H\approx0.693(R_A+R_B)C": _math(
        _variable("t", "H"),
        _mo("≈"),
        _mn("0.693"),
        _mo("("),
        _variable("R", "A"),
        _mo("+"),
        _variable("R", "B"),
        _mo(")"),
        _variable("C"),
    ),
    r"t_L\approx0.693R_BC": _math(
        _variable("t", "L"),
        _mo("≈"),
        _mn("0.693"),
        _variable("R", "B"),
        _variable("C"),
    ),
    r"f\approx\frac{1.44}{(R_A+2R_B)C}": _math(
        _variable("f"),
        _mo("≈"),
        _mfrac(
            _mn("1.44"),
            _mrow(
                _mo("("),
                _variable("R", "A"),
                _mo("+"),
                _mn("2"),
                _variable("R", "B"),
                _mo(")"),
                _variable("C"),
            ),
        ),
    ),
    r"R_P=\frac{1}{2}\left(\frac{1.44}{fC}-R_A\right)-4.7\ \mathrm{k\Omega}\approx64.95\ \mathrm{k\Omega}": _math(
        _variable("R", "P"),
        _mo("="),
        _mfrac(_mn("1"), _mn("2")),
        _mo("("),
        _mfrac(_mn("1.44"), _mrow(_variable("f"), _variable("C"))),
        _mo("−"),
        _variable("R", "A"),
        _mo(")"),
        _mo("−"),
        _mn("4.7"),
        _unit("kΩ"),
        _mo("≈"),
        _mn("64.95"),
        _unit("kΩ"),
    ),
    r"EG=S_0,\quad EY=S_1,\quad ER=\overline{EG+EY}": _math(
        _signal("EG"),
        _mo("="),
        _variable("S", "0"),
        _mo(","),
        _EM_SPACE,
        _signal("EY"),
        _mo("="),
        _variable("S", "1"),
        _mo(","),
        _EM_SPACE,
        _signal("ER"),
        _mo("="),
        _mover(_mrow(_signal("EG"), _mo("+"), _signal("EY"))),
    ),
    r"NG=S_2,\quad NY=S_3,\quad NR=\overline{NG+NY}": _math(
        _signal("NG"),
        _mo("="),
        _variable("S", "2"),
        _mo(","),
        _EM_SPACE,
        _signal("NY"),
        _mo("="),
        _variable("S", "3"),
        _mo(","),
        _EM_SPACE,
        _signal("NR"),
        _mo("="),
        _mover(_mrow(_signal("NG"), _mo("+"), _signal("NY"))),
    ),
    r"\tau=RC=150\ \mathrm{k\Omega}\times1\ \mathrm{\mu F}=0.15\ \mathrm{s}": _math(
        _variable("τ"),
        _mo("="),
        _variable("R"),
        _variable("C"),
        _mo("="),
        _mn("150"),
        _unit("kΩ"),
        _mo("×"),
        _mn("1"),
        _unit("μF"),
        _mo("="),
        _mn("0.15"),
        _unit("s"),
    ),
}


def set_run_font(run, east_asia=FONT_BODY, latin=FONT_LATIN, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR_BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(
    paragraph,
    *,
    before=0,
    after=0,
    line=20,
    first_line=True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    keep_with_next=False,
    keep_together=False,
    page_break_before=False,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = Pt(line) if line is not None else None
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    fmt.widow_control = False
    fmt.page_break_before = page_break_before
    paragraph.alignment = align


def add_page_number(paragraph, field_format, cached_value):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"PAGE \\* {field_format}"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = cached_value
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run_font(run, size=9)


def set_page_number_format(section, number_format):
    section_properties = section._sectPr
    page_number_type = section_properties.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section_properties.insert_element_before(page_number_type, "w:cols", "w:docGrid")
    page_number_type.set(qn("w:fmt"), number_format)
    page_number_type.set(qn("w:start"), "1")


def configure_section(section, number_format, field_format, cached_value, *, new_footer=False):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    section.header.paragraphs[0].text = ""
    if new_footer:
        # Remove the cloned reference before creating this section's footer part.
        section.footer.is_linked_to_previous = True
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.text = ""
    set_paragraph_format(
        footer,
        line=12,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_page_number(footer, field_format, cached_value)
    set_page_number_format(section, number_format)


def configure_document(doc):
    configure_section(doc.sections[0], "upperRoman", "ROMAN", "I")

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = COLOR_BLACK
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.widow_control = False

    configure_style(doc, "Heading 1", FONT_HEI, 12, True, 3, 3)
    configure_style(doc, "Heading 2", FONT_HEI, 12, True, 0, 0)
    configure_style(doc, "Heading 3", FONT_BODY, 12, False, 0, 0)

    add_custom_style(doc, "Table Caption", FONT_HEI, TABLE_SIZE, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_custom_style(doc, "Figure Caption", FONT_HEI, TABLE_SIZE, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_custom_style(doc, "Equation", FONT_MATH, 12, False, WD_ALIGN_PARAGRAPH.CENTER)

    doc.core_properties.title = "电子技术课程设计报告"
    doc.core_properties.subject = "电子技术课程设计"
    doc.core_properties.author = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = "电子技术课程设计, Multisim, 数字电路"


def add_toc(doc):
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph,
        line=12,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))

    update_fields = doc.settings._element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings._element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def start_body_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, "decimal", "ARABIC", "1", new_footer=True)


def configure_style(doc, name, east_asia, size, bold, before, after):
    style = doc.styles[name]
    style.font.name = FONT_LATIN
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = Pt(20)
    style.paragraph_format.left_indent = Cm(0)
    style.paragraph_format.right_indent = Cm(0)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.keep_with_next = False
    style.paragraph_format.keep_together = False
    style.paragraph_format.widow_control = False


def add_custom_style(doc, name, east_asia, size, bold, align):
    if name in doc.styles:
        style = doc.styles[name]
    else:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = FONT_LATIN if east_asia != FONT_MATH else FONT_MATH
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.alignment = align
    style.paragraph_format.left_indent = Cm(0)
    style.paragraph_format.right_indent = Cm(0)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.keep_together = True
    style.paragraph_format.widow_control = False
    return style


def clean_inline(text):
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    return text


def get_math_xslt():
    global _MATH_XSLT
    if _MATH_XSLT is None:
        if not WORD_MATH_XSL.exists():
            raise FileNotFoundError(f"Microsoft Word 公式转换文件不存在: {WORD_MATH_XSL}")
        _MATH_XSLT = etree.XSLT(etree.parse(str(WORD_MATH_XSL)))
    return _MATH_XSLT


def latex_to_omml(source, *, size=BODY_SIZE):
    source = source.strip()
    try:
        mathml_source = MATHML_EXPRESSIONS[source]
    except KeyError as exc:
        raise ValueError(f"未配置的数学表达式: {source}") from exc

    mathml = etree.fromstring(mathml_source.encode("utf-8"))
    transformed = get_math_xslt()(mathml)
    omml = transformed.getroot()
    if omml.tag != f"{{{MATH_NS}}}oMath":
        raise ValueError(f"公式转换未生成 m:oMath: {source}")

    half_points = str(int(round(size * 2)))
    for math_run in omml.xpath(".//m:r", namespaces={"m": MATH_NS}):
        word_r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), FONT_MATH)
        word_r_pr.append(fonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        word_r_pr.append(color)
        size_node = OxmlElement("w:sz")
        size_node.set(qn("w:val"), half_points)
        word_r_pr.append(size_node)
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), half_points)
        word_r_pr.append(size_cs)

        math_r_pr = math_run.find(qn("m:rPr"))
        insert_at = 1 if math_r_pr is not None else 0
        math_run.insert(insert_at, word_r_pr)
    return omml


def add_inline_content(paragraph, text, *, east_asia=FONT_BODY, latin=FONT_LATIN, size=BODY_SIZE, bold=None):
    text = clean_inline(text)
    position = 0
    for match in re.finditer(r"\$([^$]+)\$", text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)
        paragraph._p.append(latex_to_omml(match.group(1), size=size))
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def add_heading(doc, text, level):
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(clean_inline(text))
    if level <= 2:
        set_run_font(run, east_asia=FONT_HEI, latin=FONT_LATIN, size=12, bold=True)
    else:
        set_run_font(run, east_asia=FONT_BODY, latin=FONT_LATIN, size=12, bold=False)


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    add_inline_content(p, text)
    return p


def add_equation(doc, source, number):
    p = doc.add_paragraph(style="Equation")
    set_paragraph_format(
        p,
        before=3,
        after=3,
        line=None,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        keep_together=True,
    )
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    lead_tab = p.add_run("\t")
    set_run_font(lead_tab, size=BODY_SIZE)
    p._p.append(latex_to_omml(source, size=BODY_SIZE))
    number_run = p.add_run(f"\t({number})")
    set_run_font(number_run, size=BODY_SIZE)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_border_edge(cell, edge, *, value="single", color="000000", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    tag = qn(f"w:{edge}")
    element = borders.find(tag)
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), value)
    element.set(qn("w:sz"), size if value != "nil" else "0")
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)


def apply_three_line_borders(table):
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right"):
                set_cell_border_edge(cell, edge, value="nil")
    for cell in table.rows[0].cells:
        set_cell_border_edge(cell, "top", size="12")
        set_cell_border_edge(cell, "bottom", size="6")
    for cell in table.rows[-1].cells:
        set_cell_border_edge(cell, "bottom", size="12")


def set_cell_margins(cell, top=0, start=108, bottom=0, end=108):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def table_widths(column_count):
    if column_count == 2:
        return [5.5, 10.5]
    if column_count == 3:
        return [2.8, 7.8, 5.4]
    if column_count == 6:
        return [2.0, 2.6, 2.6, 2.1, 2.1, 4.6]
    width = 16.0 / column_count
    return [width] * column_count


def cm_to_dxa(value):
    return int(value / 2.54 * 1440)


def set_table_geometry(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths_dxa = [cm_to_dxa(value) for value in widths_cm]
    total_dxa = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height = Pt(16.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            set_cell_margins(cell)
            set_cell_shading(cell)


def add_table_caption(doc, text):
    p = doc.add_paragraph(style="Table Caption")
    set_paragraph_format(
        p,
        before=0,
        after=6,
        line=20,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    run = p.add_run(clean_inline(text))
    set_run_font(run, east_asia=FONT_HEI, latin=FONT_HEI, size=TABLE_SIZE, bold=False)


def add_table(doc, rows):
    column_count = len(rows[0])
    widths = table_widths(column_count)
    table = doc.add_table(rows=1, cols=column_count)

    for row_index, row_data in enumerate(rows):
        row = table.rows[0] if row_index == 0 else table.add_row()
        for col_index, value in enumerate(row_data):
            cell = row.cells[col_index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_format(
                p,
                line=None,
                first_line=False,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=False,
            )
            p.paragraph_format.line_spacing = 1.0
            add_inline_content(p, value, size=TABLE_SIZE, bold=False)
    set_table_geometry(table, widths)
    apply_three_line_borders(table)
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, after=6, line=20, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_image(doc, image_path, alt_text):
    width_map = {
        "fig01": 15.5,
        "fig02": 15.5,
        "fig03": 15.0,
        "fig04": 15.0,
        "fig05": 15.0,
        "fig06": 14.5,
    }
    width = 15.8
    for prefix, value in width_map.items():
        if image_path.name.startswith(prefix):
            width = value
            break
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        before=0,
        after=6,
        line=None,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
        keep_together=True,
    )
    # Inline pictures must use automatic line height.  Otherwise Word inherits
    # the Normal style's fixed 20 pt line spacing and clips the picture.
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width))
    return alt_text


def add_figure_caption(doc, text):
    p = doc.add_paragraph(style="Figure Caption")
    set_paragraph_format(
        p,
        before=0,
        after=6,
        line=20,
        first_line=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        keep_together=True,
    )
    run = p.add_run(clean_inline(text))
    set_run_font(run, east_asia=FONT_HEI, latin=FONT_HEI, size=TABLE_SIZE, bold=False)


def parse_table(lines, start):
    raw_rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        raw_rows.append(cells)
        index += 1
    rows = []
    for row in raw_rows:
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in row):
            continue
        rows.append(row)
    return rows, index


def build():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_toc(doc)
    start_body_section(doc)

    index = 0
    equation_number = 1
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        if stripped == "$$":
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation_lines.append(lines[index].strip())
                index += 1
            add_equation(doc, " ".join(equation_lines), equation_number)
            equation_number += 1
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            image_path = (SOURCE_MD.parent / image_match.group(2)).resolve()
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            add_image(doc, image_path, image_match.group(1))
            index += 1
            continue

        if stripped.startswith("# "):
            # The source H1 is document metadata; the report starts with section 1.
            index += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            index += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            index += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            index += 1
            continue

        if stripped.startswith("**表") and stripped.endswith("**"):
            add_table_caption(doc, stripped[2:-2].strip())
            index += 1
            continue
        if stripped.startswith("**图") and stripped.endswith("**"):
            add_figure_caption(doc, stripped[2:-2].strip())
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        add_body(doc, stripped)
        index += 1

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
